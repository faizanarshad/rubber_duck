"""Generate a .docx report for the RubberDuck security-sensitive path audit."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


OUTPUT_PATH = Path(__file__).resolve().parent.parent / "security_audit_rubberduck.docx"


def set_cell_background(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def risk_color(level: str) -> str:
    return {
        "High": "F8D7DA",
        "Medium": "FFF3CD",
        "Low": "D4EDDA",
        "None": "E2E3E5",
    }.get(level, "FFFFFF")


def build_table(doc: Document, headers, rows, col_widths_cm=None, risk_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "1F3A68")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if risk_col is not None and c_idx == risk_col:
                set_cell_background(cell, risk_color(str(val)))
                run.bold = True

    if col_widths_cm:
        for row in table.rows:
            for i, width in enumerate(col_widths_cm):
                row.cells[i].width = Cm(width)

    return table


def build_report() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Security-Sensitive Path Audit")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("RubberDuck Semantic Intelligence — rubber_duck/backend")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   Audit duration: ~100 s"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    add_heading(doc, "Executive Summary", level=1)
    add_paragraph(
        doc,
        "A full data-flow audit of the loaded backend (13 source files) was performed using the "
        "RubberDuck semantic-intelligence MCP (load_code, symbols_overview, search_code, "
        "trace_variable, call_chain). No high- or medium-risk dangerous-sink paths reachable from "
        "external input were found. The only user-reachable sinks are file operations on OS-generated "
        "temporary paths derived from a closed extension allowlist ('.pdf', '.csv'). Zero occurrences "
        "of exec / eval / subprocess / os.system / shell=True / pickle.loads / yaml.load / raw SQL / "
        "outbound HTTP libraries were found in the backend scope.",
    )

    add_heading(doc, "Scope of Analysis", level=1)
    add_paragraph(doc, "Loaded backend modules (analysis IDs in parentheses):")
    for line in [
        "• backend/main.py (main)",
        "• backend/api/routes_chat.py (routes_chat)",
        "• backend/api/routes_files.py (routes_files)",
        "• backend/core/config.py (config)",
        "• backend/medical_config.py (medical_config)",
        "• backend/services/rag_service.py (rag_service)",
        "• backend/services/vectordb_service.py (vectordb_service)",
        "• backend/services/embeddings_service.py (embeddings_service)",
        "• backend/services/llm_service.py (llm_service)",
        "• backend/services/data_injestion_service.py (data_injestion_service)",
        "• backend/services/csv_processor.py (csv_processor)",
        "• backend/utils/logger.py (logger)",
        "• backend/test_csv.py (test_csv)",
    ]:
        add_paragraph(doc, line, size=10)
    add_paragraph(
        doc,
        "Third-party PyPDF2 files that were also loaded were excluded from the audit (library code, "
        "not part of this codebase).",
        italic=True,
        size=10,
    )

    add_heading(doc, "Step 1 — Entry Points Accepting External Input", level=1)
    add_paragraph(
        doc,
        "Found via search_code with pattern: "
        "@router\\.(get|post|put|delete|patch)|@app\\.(get|post|put|delete|patch)",
        italic=True,
        size=10,
    )
    build_table(
        doc,
        headers=["#", "Entry point (method + route)", "File : Line", "External input"],
        rows=[
            ["1", "POST  chat  /",                 "routes_chat.py : 35–81",   "request: ChatRequest (body: query: str)"],
            ["2", "POST  add_file  /add_file",     "routes_files.py : 127–218", "file: UploadFile"],
            ["3", "DEL   delete_file /delete_file/{file_id}", "routes_files.py : 221–263", "file_id: str (path)"],
            ["4", "PUT   update_file /update_file/{file_id}", "routes_files.py : 266–341", "file_id: str, file: UploadFile"],
            ["5", "POST  get_csv_info  /csv_info", "routes_files.py : 344–409", "file: UploadFile"],
            ["6", "GET   health / health_check",   "main.py:76 / routes_*.py",  "(no user input)"],
            ["7", "GET   root  /",                 "main.py : 60",              "(no user input)"],
        ],
        col_widths_cm=[1.0, 5.5, 4.5, 5.5],
    )

    add_heading(doc, "Step 2 — Data Flow of Each Input (trace_variable)", level=1)
    flow_rows = [
        ["request (chat body)",
         "4 nodes in chat (lines 36, 50, 53, 61). Reaches request.query.strip() and rag_service.process_query(request.query) on line 61. No cross-module file / shell / SQL flow."],
        ["file (UploadFile)",
         "14 nodes across add_file, update_file, get_csv_info. Reaches only file.filename.lower().endswith(...), file.read(), and logger.info. Never concatenated into a path, command, or query."],
        ["file.filename",
         "Used at lines 142, 148, 189, 197, 285, 295, 356, 388 — log messages, allowlist check, and as metadata in process_csv_file. No path construction."],
        ["file_id (path param)",
         "15 nodes. Reaches rag_service.delete_document / update_document. Downstream becomes a Pinecone metadata filter {'file_id': {'$eq': file_id}} and a vector-id prefix f'{file_id}_{i}'. No shell / SQL / FS sink."],
        ["suffix → NamedTemporaryFile",
         "Assigned from file_ext at line 159. file_ext is set only inside an allowlist loop over ['.pdf', '.csv'] (lines 145–150); otherwise None → HTTP 400. Not user-controlled."],
        ["temp_file_path",
         "Assigned from temp_file.name — OS-generated by tempfile.NamedTemporaryFile. User bytes never influence the path string."],
    ]
    build_table(
        doc,
        headers=["Input variable", "Traced flow"],
        rows=flow_rows,
        col_widths_cm=[5.0, 11.5],
    )

    add_heading(doc, "Step 3 — Dangerous Sink Search (search_code)", level=1)
    add_paragraph(
        doc,
        "All regexes were run across all 13 loaded backend analyses.",
        italic=True,
        size=10,
    )
    sink_rows = [
        ["exec / eval",                              r"\b(exec|eval)\s*\(",                                                    "0"],
        ["subprocess / os.system / Popen / shell=True", "subprocess|os\\.system|os\\.popen|shell=True|check_output|Popen",    "0"],
        ["pickle / marshal / yaml.load / __import__ / compile()", "pickle\\.|marshal\\.|yaml\\.load|__import__|compile\\s*\\(", "0 in backend (1 LangGraph workflow.compile() — not Python builtin compile)"],
        ["SQL (execute / cursor.execute / SELECT / INSERT / UPDATE / DELETE FROM / raw( / text( / .query()",
         "execute\\s*\\(|cursor\\.execute|executemany|raw\\s*\\(|text\\s*\\(|\\.query\\s*\\(|SELECT |INSERT |UPDATE |DELETE FROM",
         "0 SQL. The .query() hits are Pinecone self.index.query(...) at vectordb_service.py:114, 157 — vector search, not SQL."],
        ["HTTP out (requests / urllib / urlopen / httpx)", "requests\\.|urllib|urlopen|httpx\\.",                              "0 functional calls (only a comment in config.py:10)"],
        ["Template injection (jinja2 / Template( / format_map / render_template_string)",
         "render_template_string|jinja2|Template\\(|format_map",
         "0 (one static f-string template in medical_config.get_medical_prompt_template)"],
        ["File ops (open( / .read( / .write( / NamedTemporaryFile / os.unlink / shutil / Path()",
         "open\\s*\\(|\\.write\\s*\\(|\\.read\\s*\\(|NamedTemporaryFile|os\\.remove|os\\.unlink|shutil\\.|Path\\(",
         "Present — see Step 4"],
    ]
    build_table(
        doc,
        headers=["Sink category", "Regex used", "Backend matches"],
        rows=sink_rows,
        col_widths_cm=[4.5, 5.5, 6.5],
    )

    add_heading(doc, "Step 4 — Reachability From User Input to Existing Sinks (call_chain)", level=1)
    add_paragraph(
        doc,
        "call_chain was run on every entry-point handler and every transitive service method "
        "(chat, add_file, process_csv_file, update_file, delete_file, process_query, retrieve_documents, "
        "generate_answer, add_document, update_document, delete_document, search_similar, "
        "process_pdf_file, process_csv_to_documents, generate_response, process_documents).",
        italic=True,
        size=10,
    )
    reach_rows = [
        ["tempfile.NamedTemporaryFile(delete=False, suffix=suffix)",
         "routes_files.py : 160",
         "Temp-file create",
         "Yes — suffix ∈ {'.pdf', '.csv'} from allowlist (lines 145–159). Not attacker-controlled.",
         "Low"],
        ["tempfile.NamedTemporaryFile(... suffix='.pdf')",
         "routes_files.py : 302",
         "Temp-file create",
         "Hard-coded suffix.",
         "Low"],
        ["tempfile.NamedTemporaryFile(... suffix='.csv')",
         "routes_files.py : 363",
         "Temp-file create",
         "Hard-coded suffix.",
         "Low"],
        ["temp_file.write(await file.read())",
         "routes_files.py : 162, 304, 365",
         "Write uploaded bytes",
         "Content is user bytes, path is OS-generated. No traversal.",
         "Low"],
        ["os.unlink(temp_file_path)",
         "routes_files.py : 209, 332, 402",
         "Delete temp file",
         "Path from temp_file.name (OS). Guarded by os.path.exists.",
         "Low"],
        ["open(pdf_file_path, 'rb')",
         "data_injestion_service.py : 41, 135",
         "File read (PDF parse)",
         "Argument is temp_file_path (tempfile), not file.filename.",
         "Low"],
        ["Path(file_path).stat()",
         "csv_processor.py : 418",
         "File stat",
         "Same — OS tempfile path.",
         "Low"],
    ]
    build_table(
        doc,
        headers=["Sink (code)", "File : Line", "Type", "Input reaches it?", "Risk"],
        rows=reach_rows,
        col_widths_cm=[5.0, 3.3, 2.8, 4.2, 1.2],
        risk_col=4,
    )

    add_heading(doc, "Final Findings", level=1)
    final_rows = [
        ["1", "add_file",
         "file.filename → allowlist .endswith(ext) → file_ext ∈ {'.pdf','.csv'} → suffix → NamedTemporaryFile(suffix=suffix)",
         "routes_files.py : 160",
         "Temp-file create with allowlisted literal suffix",
         "Low"],
        ["2", "add_file / update_file / get_csv_info",
         "await file.read() → temp_file.write(content) at OS-generated path",
         "routes_files.py : 162, 304, 365",
         "Write user bytes to tempfile",
         "Low"],
        ["3", "add_file / update_file / get_csv_info",
         "temp_file.name → temp_file_path → os.unlink(...) (guarded by os.path.exists)",
         "routes_files.py : 209, 332, 402",
         "Delete OS-generated temp path",
         "Low"],
        ["4", "add_file → rag_service.add_document → DataIngestionService.process_pdf_file",
         "temp_file_path → open(pdf_file_path, 'rb')",
         "data_injestion_service.py : 41, 135",
         "File read on OS-generated path",
         "Low"],
        ["5", "delete_file / update_file",
         "file_id → rag_service.delete_document / update_document → Pinecone filter {'file_id': {'$eq': file_id}} and vector-id f'{file_id}_{i}'",
         "vectordb_service.py : 143–177; rag_service.py : 234",
         "Vector-DB metadata filter (no SQL)",
         "Low"],
    ]
    build_table(
        doc,
        headers=["#", "Entry point", "Data flow path", "Sink", "Type", "Risk"],
        rows=final_rows,
        col_widths_cm=[0.9, 3.3, 5.5, 3.3, 2.5, 1.2],
        risk_col=5,
    )

    add_heading(doc, "Honest Conclusion", level=1)
    add_paragraph(
        doc,
        "No high- or medium-risk dangerous-sink paths were found in the loaded backend scope. "
        "There are zero occurrences of exec, eval, subprocess, os.system, shell=True, Popen, "
        "pickle.loads, yaml.load, __import__, raw SQL, cursor.execute, or outbound HTTP libraries "
        "across any of the 13 backend files. The only user-reachable sinks are tempfile create / "
        "write / unlink and open() on OS-generated paths, all with either a hard-coded suffix or a "
        "two-element allowlist ('.pdf', '.csv') — none of them let an attacker control the filesystem path.",
    )

    add_heading(doc, "Out-of-Scope Observations", level=1)
    add_paragraph(
        doc,
        "The following are not data-flow sinks but are worth noting. They were not counted in the "
        "risk table.",
        italic=True,
        size=10,
    )
    obs_rows = [
        ["CORS misconfiguration",
         "main.py : 47–53",
         "CORSMiddleware(allow_origins=['*'], allow_credentials=True). FastAPI rejects this combination at runtime; also weakens origin checks.",
         "Medium"],
        ["Error detail leakage",
         "routes_files.py : 214, 217, 259, 262, 337, 340; routes_chat.py : 77",
         "HTTPException(detail=f\"… {str(e)}\") echoes raw exception messages to clients.",
         "Low"],
        ["Prompt-injection surface (LLM, not OS)",
         "routes_chat.py → rag_service.process_query → llm_service.generate_response : 47",
         "Attacker-controlled request.query reaches openai.chat.completions.create. Classic prompt injection risk, not an OS / SQL / shell sink.",
         "Medium"],
    ]
    build_table(
        doc,
        headers=["Observation", "Location", "Description", "Risk"],
        rows=obs_rows,
        col_widths_cm=[3.5, 4.5, 7.3, 1.2],
        risk_col=3,
    )

    add_heading(doc, "Methodology", level=1)
    methodology_items = [
        "1. load_code(repo='local/rubber_duck', subpath='backend') — loaded 13 backend source files into the semantic server.",
        "2. search_code — located HTTP decorators (@router.*, @app.*) to enumerate entry points.",
        "3. search_code — ran 7 regex patterns for dangerous sinks (exec/eval, subprocess, pickle/marshal/compile, SQL, HTTP-out, template injection, file ops).",
        "4. query_action(action='trace_variable') — traced each user input variable (request, file, file_id, file.filename, suffix, file_ext, temp_file_path) across all scopes.",
        "5. query_action(action='call_chain') — traced every entry-point handler and every downstream service method to confirm whether any tainted input can reach a dangerous sink.",
        "6. Read-verified line 145–150 of routes_files.py to confirm the allowlist loop semantics.",
    ]
    for item in methodology_items:
        add_paragraph(doc, item, size=10)

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "Audit duration (wall clock, measured via date +%s before and after): 1776461919 − 1776461819 = 100 seconds.",
        italic=True,
        size=10,
    )

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
