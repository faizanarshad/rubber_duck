#!/usr/bin/env python3
"""Generate audit_results_compiled.docx with consolidated analysis results."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "audit_results_compiled.docx"

    doc = Document()
    sect = doc.sections[0]
    sect.page_height = sect.page_height  # noqa: B018 — keep default

    title = doc.add_heading("Rubber Duck / Agentic RAG — Consolidated Audit Results", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_p(
        doc,
        "Generated document compiling security traces, tooling outcomes, medical/HIPAA "
        "gap analysis, and methodology comparison. Domain: healthcare-oriented document RAG "
        "(FastAPI, React, OpenAI, Pinecone, PDF/CSV ingestion).",
    )
    add_p(doc, "Compilation date: April 16, 2026.")

    add_h(doc, "1. Repository & run verification", 1)
    add_bullets(
        doc,
        [
            "Remote: origin/main (faizanarshad/rubber_duck); git pull reported already up to date.",
            "Backend: import main OK; GET /health returned healthy JSON when server started locally.",
            "Frontend: npm run build (tsc + vite) succeeded.",
        ],
    )

    add_h(doc, "2. MCP configuration (summary)", 1)
    add_p(
        doc,
        "Cursor MCP (~/.cursor/mcp.json) configured for rubberduck-codebase-intelligence "
        "and rubberduck-semantic-intelligence (HTTPS endpoints; Authorization: Bearer token — "
        "not reproduced in this document). Restart Cursor after changes to reload MCP servers.",
    )

    add_h(doc, "3. Semantic intelligence: load_code", 1)
    add_bullets(
        doc,
        [
            "Repo loaded: faizanarshad/rubber_duck (max_files=100).",
            "Errors: 0; total Python files reported in snapshot: 6360 (includes vendored paths on server).",
            "Sample analysis_ids: routes_chat, routes_files, main, config, rag_service, vectordb_service, etc.",
        ],
    )

    add_h(doc, "4. Codebase-intelligence security tooling", 1)
    add_bullets(
        doc,
        [
            "get_security_report: not exposed as a named MCP tool in this Cursor environment.",
            "assess(mode=security) with OWASP LLM Top 10 scope: verdict clean, 0 findings, 50 files; "
            "CI findings unavailable; codebase_intelligence fusion false.",
            "assess(mode=repo_quality): cached metrics for faizanarshad/rubber_duck (e.g. overall health ~89.5, grade A) — quality, not HIPAA.",
            "discover/list_repos + index list: repos known but indexing status NOT_INDEXED at time of check.",
        ],
    )

    add_h(doc, "5. Path audit (loaded analyzer/web scope — not main backend RAG)", 1)
    add_p(
        doc,
        "Note: This pass targeted FastAPI modules loaded in Rubber Duck (e.g. analysis.py, jobs.py, "
        "sessions.py, synthesis.py), not necessarily the Agentic RAG backend/ tree.",
    )
    add_bullets(
        doc,
        [
            "Entry points: APIRouter handlers (upload, job status, sessions, synthesis reports).",
            "High risk (example): UploadFile.filename → _save_upload → dest.write_bytes — path traversal / "
            "unsafe filename if unsanitized.",
            "Lower risk (guarded): session delete via _safe_session_path; synthesis delete_report with "
            "filename/path checks.",
            "No exec/eval/subprocess/SQL execute sinks identified in that loaded scope.",
        ],
    )

    add_h(doc, "6. Entry-to-sink traces (backend/ Agentic RAG)", 1)

    add_h(doc, "6a. File upload → CSV → Pinecone", 2)
    add_bullets(
        doc,
        [
            "POST /files/add_file → temp file → CSV branch → CSVProcessor.process_csv_to_documents.",
            "EmbeddingsService.process_documents → VectorDBService.upsert_documents → Pinecone upsert; "
            "metadata includes full chunk text.",
            "Risks: no app-layer auth in snippets; large CSV DoS; PHI in cell text if anonymization misses.",
        ],
    )

    add_h(doc, "6b. User query → vector DB → LLM", 2)
    add_bullets(
        doc,
        [
            "POST /chat/ → RAGService.process_query → retrieve: generate_embedding(query) → "
            "search_similar → generate: LLMService.generate_response.",
            "Prompt combines retrieved context + user query; sent to OpenAI chat.completions.",
            "Risks: prompt injection (query + retrieved text); PHI in prompts; ChatResponse may return "
            "context_documents with full retrieved text.",
        ],
    )

    add_h(doc, "7. Medical / HIPAA — gap summary", 1)
    add_bullets(
        doc,
        [
            "MedicalRAGConfig.HIPAA_COMPLIANCE flags (audit, access, encryption, auth) are metadata only — "
            "not enforced in FastAPI routes.",
            "LLMService uses generic assistant system prompt; MedicalRAGConfig medical prompts not wired to live chat.",
            "CSV anonymize_data(): column-name heuristics + weak name regex; misses PHI in free-text columns.",
            "PDF ingestion: no PHI scrubbing — raw extract → chunk → embed (high risk for clinical PDFs).",
            "Logging may log query prefixes — PHI if users paste clinical text.",
            "Subprocessors: full text to OpenAI embeddings/chat; full text in Pinecone metadata — requires "
            "BAA/risk analysis for ePHI use cases.",
        ],
    )

    add_h(doc, "8. Comparison: audit methodologies", 1)
    add_p(doc, "What each approach found vs missed (abbreviated):")

    table = doc.add_table(rows=5, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Approach"
    hdr[1].text = "Strong at"
    hdr[2].text = "Weak / missed"
    hdr[3].text = "Medical-domain note"

    rows_data = [
        (
            "Semantic graph audit",
            "In-file flows to sinks in loaded code",
            "Unloaded paths; policy/HIPAA",
            "Did not validate PHI via OpenAI/Pinecone",
        ),
        (
            "assess / CI codebase intel",
            "Repo metrics when indexed",
            "Tool gaps; unfused security run",
            "Clean ≠ HIPAA-safe for clinical data",
        ),
        (
            "Manual E2E trace",
            "Actual RAG data flow in backend/",
            "No exploit proof",
            "context_documents exposes retrieved text",
        ),
        (
            "Medical code review",
            "Config vs code; PDF/CSV PHI",
            "Legal BAAs, workforce",
            "PDF raw text → embeddings is key clinical risk",
        ),
    ]
    for i, row in enumerate(rows_data, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    add_p(doc, "")
    add_p(
        doc,
        "Ship recommendation: Do not rely on graph clean or assess clean alone for PHI/medical deployment. "
        "Require manual trace + medical gaps addressed, plus auth, minimum-necessary API, de-ID strategy, "
        "and subprocessor alignment.",
    )

    add_h(doc, "9. Files referenced", 1)
    add_bullets(
        doc,
        [
            "backend/api/routes_files.py, routes_chat.py",
            "backend/services/rag_service.py, llm_service.py, vectordb_service.py, csv_processor.py, data_injestion_service.py",
            "backend/medical_config.py",
            "~/.cursor/mcp.json (MCP servers)",
        ],
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
